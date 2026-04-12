"""PDF<->Word conversion. PDF->Word uses PyMuPDF+python-docx. Word->PDF uses LibreOffice CLI."""
import fitz
import subprocess
import shutil
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)


def pdf_to_word(input_path: Path, output_path: Path) -> Path:
      doc_out = Document()
      for section in doc_out.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            with fitz.open(str(input_path)) as pdf:
                      for page_num, page in enumerate(pdf):
                                    if page_num > 0:
                                                      doc_out.add_page_break()
                                                  blocks = page.get_text("dict")["blocks"]
                                    for block in blocks:
                                                      if block.get("type") != 0:
                                                                            continue
                                                                        for line in block.get("lines", []):
                                                                                              line_text = ""
                                                                                              max_size = 11
                                                                                              is_bold = False
                                                                                              for span in line.get("spans", []):
                                                                                                                        line_text += span.get("text", "")
                                                                                                                        s = span.get("size", 11)
                                                                                                                        if s > max_size:
                                                                                                                                                      max_size = s
                                                                                                                                                  if span.get("flags", 0) & (1 << 4):
                                                                                                                                                                                is_bold = True
                                                                                                                                                                        if line_text.strip():
                                                                                                                                                                                                  para = doc_out.add_paragraph()
                                                                                                                                                                                                  run = para.add_run(line_text)
                                                                                                                                                                                                  run.font.size = Pt(max_size)
                                                                                                                                                                                                  run.bold = is_bold
                                                                                                                                                                              doc_out.save(str(output_path))
                                                                                                    return output_path


                        def word_to_pdf(input_path: Path, output_path: Path) -> Path:
                              lo = _find_libreoffice()
                              if not lo:
                                        raise RuntimeError("LibreOffice not found. Install: sudo apt-get install -y libreoffice")
                                    return _lo_convert(lo, input_path, output_path.parent, output_path)


def _find_libreoffice():
      for c in ["libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice",
                              "/usr/lib/libreoffice/program/soffice"]:
                                        if shutil.which(c):
                                                      return c
                                              return None


def _lo_convert(lo, input_path, output_dir, final_path):
      result = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(input_path)],
                capture_output=True, text=True, timeout=120
      )
    if result.returncode != 0:
              raise RuntimeError(f"LibreOffice error: {result.stderr or result.stdout}")
    converted = output_dir / f"{input_path.stem}.pdf"
    if not converted.exists():
              raise RuntimeError("LibreOffice did not produce output file.")
    if converted != final_path:
              converted.rename(final_path)
    return final_path
